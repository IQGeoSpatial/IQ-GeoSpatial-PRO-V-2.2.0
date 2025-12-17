import os
import tempfile
import shutil
import pandas as pd
import re
import unicodedata
import re as _re

from docx.shared import Cm
from docx import Document
from docx.table import _Row
from utils.resource_path import resource_path
from copy import deepcopy
from docx.shared import Inches


# --- FUNCIÓN DE DEPURACIÓN SOLO PARA FORMULARIO 001 ---
def debug_imprime_contexto_formulario_001(excel_path):
    """
    Lee la hoja 1 del Excel y muestra en consola las claves y valores que se usarían para el Formulario 001.
    """
    def normaliza_clave(clave):
        clave = str(clave).strip()
        clave = unicodedata.normalize('NFKD', clave).encode('ASCII', 'ignore').decode('ASCII')
        clave = clave.replace(' ', '_').replace('-', '_').replace('__', '_').upper()
        return clave

    df = pd.read_excel(excel_path, sheet_name=0, header=None)
    # print('--- CONTENIDO DE LA HOJA 1 (Formulario 001) ---')
    for i in range(len(df)):
        nombre = str(df.iloc[i, 0]).strip()
        valor = df.iloc[i, 1] if df.shape[1] > 1 else ''
        if nombre != '':
            valor_str = str(valor).strip() if pd.notna(valor) else ''
            # print(f"CLAVE ORIGINAL: '{nombre}' | NORMALIZADA: '{normaliza_clave(nombre)}' | VALOR: '{valor_str}'")


class WordGenerator:
    def _contexto_formulario_001(self):
        """Construye el contexto para Formulario 001 usando hoja 1, columna A como nombre y columna B como valor, agregando clave original y todas las variantes normalizadas. (NO TRANSPONE)"""
        contexto = {}
        def normaliza_clave(clave):
            clave = str(clave).strip()
            clave = unicodedata.normalize('NFKD', clave).encode('ASCII', 'ignore').decode('ASCII')
            clave = clave.replace(' ', '_').replace('-', '_').replace('__', '_').upper()
            return clave
        # Leer hoja 1 directamente, sin transponer
        try:
            df = pd.read_excel(self.excel_path, sheet_name=0, header=None)
        except Exception as e:
            raise ValueError(f"Error al leer la hoja 1 del Excel para Formulario 001: {e}")
        for i in range(len(df)):
            nombre = str(df.iloc[i, 0]).strip()
            valor = df.iloc[i, 1] if df.shape[1] > 1 else ''
            if nombre != '':
                valor_str = str(valor).strip() if pd.notna(valor) else ''
                claves = set()
                claves.add(nombre)
                claves.add(nombre.upper())
                claves.add(normaliza_clave(nombre))
                claves.add(normaliza_clave(nombre.lower()))
                claves.add(normaliza_clave(nombre.upper()))
                for k in claves:
                    contexto[k] = valor_str
        return contexto
    """
    Servicio para generar documentos de Word a partir de plantillas y datos de Excel.
    """
    def __init__(self, excel_path, output_dir, seleccionados):
        """Inicializa el generador, cargando los datos necesarios del Excel."""
        self.excel_path = excel_path
        self.output_dir = output_dir
        self.seleccionados = seleccionados
        self.datos_generales = {}
        self.puntos_geodesicos = pd.DataFrame()
        
        try:
            self._load_all_excel_data()
        except Exception as e:
            raise ValueError(f"Error al leer el archivo Excel: {e}")

    def _load_excel_sheet(self, sheet_index):
        """Carga y procesa una hoja específica del archivo Excel."""
        try:
            df_raw = pd.read_excel(self.excel_path, sheet_name=sheet_index, header=None)
            df_transposed = df_raw.T
            df = pd.DataFrame(df_transposed.values[1:], columns=df_transposed.iloc[0])
            df.columns = [str(col).strip() for col in df.columns]
            return df
        except Exception as e:
            raise ValueError(f"No se pudo leer la hoja {sheet_index + 1} del Excel. Verifique que el archivo no esté dañado y que la hoja exista. Error: {e}")

    def _load_all_excel_data(self):
        """Carga todos los datos necesarios del Excel una sola vez."""
        # Hoja 1: Formulario 001
        self.df_form1 = self._load_excel_sheet(0)
        # Hoja 2: Formulario 002 y 003
        self.df_form2_3 = self._load_excel_sheet(1)
        # Hoja 3: Puntos geodésicos (para 004 y 005)
        self.puntos_geodesicos = self._load_excel_sheet(2)
        if 'COD_PUNTO _GEODESICO' in self.puntos_geodesicos.columns:
            self.puntos_geodesicos.rename(columns={'COD_PUNTO _GEODESICO': 'PUNTO'}, inplace=True)
        if 'PUNTO' not in self.puntos_geodesicos.columns:
            raise ValueError("La hoja de puntos geodésicos (Hoja 3) debe contener una columna 'PUNTO'.")
        # Tomar solo filas con valor válido en PUNTO (no vacías, no nulas, no espacios)
        self.puntos_geodesicos = self.puntos_geodesicos[self.puntos_geodesicos['PUNTO'].notna() & (self.puntos_geodesicos['PUNTO'].astype(str).str.strip() != '')]

    def _safe_get(self, row, col):
        """
        Obtiene un valor de una fila de forma segura, manejando columnas duplicadas (que devuelven una Serie).
        """
        if col not in row.index:
            return ""
        val = row.get(col)

        # Si el valor es una Serie (ocurre con columnas duplicadas), tomar el primer valor no nulo.
        if isinstance(val, pd.Series):
            val = val.dropna().iloc[0] if not val.dropna().empty else ""

        # Ahora 'val' es un escalar. Se puede usar pd.isna() de forma segura.
        if pd.isna(val) or str(val).strip() == '':
            return ""
        return str(val)

    def _replace_placeholders_in_paragraph(self, paragraph, context):
        """
        Reemplaza placeholders aunque estén partidos en varios runs, usando contexto completamente normalizado.
        Conserva el formato de los runs originales (negrita, fuente, tamaño, etc.) en la medida de lo posible.
        """

        def normaliza_clave(clave):
            clave = str(clave).strip()
            clave = unicodedata.normalize('NFKD', clave).encode('ASCII', 'ignore').decode('ASCII')
            clave = clave.replace(' ', '_').replace('-', '_').replace('__', '_').upper()
            return clave

        # Reconstruir el texto completo del párrafo
        full_text = ''.join(run.text for run in paragraph.runs)
        if '{' not in full_text:
            return

        pattern = r'\{\{([A-Za-z0-9_\- \(\)]+)\}\}'
        pattern2 = r'\{([A-Za-z0-9_\- \(\)]+)\}'

        def get_replacement(key):
            key_norm = normaliza_clave(key)
            key_sin_paren = _re.sub(r'[\(\)]', '', key)
            key_sin_guion = key.replace('-', '_')
            key_sin_espacios = key.replace(' ', '')
            variantes = [
                key,
                key_norm,
                key.upper(),
                normaliza_clave(key.upper()),
                normaliza_clave(key.lower()),
                _re.sub(r'[\(\)]', '', key),
                normaliza_clave(_re.sub(r'[\(\)]', '', key)),
                key_sin_guion,
                normaliza_clave(key_sin_guion),
                key_sin_espacios,
                normaliza_clave(key_sin_espacios)
            ]
            for variante in variantes:
                if variante in context:
                    return str(context[variante])
            return None

        # Reemplazar todos los placeholders en el texto completo
        def repl1(match):
            key = match.group(1).strip()
            val = get_replacement(key)
            return val if val is not None else match.group(0)
        new_text = re.sub(pattern, repl1, full_text)
        def repl2(match):
            key = match.group(1).strip()
            val = get_replacement(key)
            return val if val is not None else match.group(0)
        new_text = re.sub(pattern2, repl2, new_text)

        # Si no hubo cambios, salir
        if new_text == full_text:
            return

        # Eliminar todos los runs originales
        for run in paragraph.runs[:]:
            p_elem = run._element
            p_elem.getparent().remove(p_elem)

        # Crear un solo run con el texto reemplazado, copiando el formato del primer run si existe
        ref_run = None
        if hasattr(paragraph, 'runs') and len(paragraph.runs) > 0:
            ref_run = paragraph.runs[0]
        new_run = paragraph.add_run(new_text)
        if ref_run:
            # Copiar formato de fuente y parrafo
            if ref_run.bold is not None:
                new_run.bold = ref_run.bold
            if ref_run.italic is not None:
                new_run.italic = ref_run.italic
            if ref_run.underline is not None:
                new_run.underline = ref_run.underline
            if ref_run.font is not None:
                if ref_run.font.name:
                    new_run.font.name = ref_run.font.name
                if ref_run.font.size:
                    new_run.font.size = ref_run.font.size
                if ref_run.font.color and ref_run.font.color.rgb:
                    new_run.font.color.rgb = ref_run.font.color.rgb
            # Copiar formato de alineación si aplica
            if hasattr(ref_run, 'alignment') and ref_run.alignment:
                new_run.alignment = ref_run.alignment

    def _replace_all_placeholders(self, doc, context):
        """
        Orquesta el reemplazo de placeholders en todo el documento, incluyendo tablas anidadas.
        """
        # Procesar párrafos en el cuerpo principal del documento
        for p in doc.paragraphs:
            self._replace_placeholders_in_paragraph(p, context)
        
        # Procesar todas las tablas, incluidas las anidadas
        for table in doc.tables:
            self._process_table_recursively(table, context)
            
        return doc

    def _process_table_recursively(self, table, context):
        """
        Procesa una tabla y cualquier tabla anidada dentro de sus celdas.
        """
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self._replace_placeholders_in_paragraph(p, context)
                for nested_table in cell.tables:
                    self._process_table_recursively(nested_table, context)

    def _fill_geodetic_points_table(self, doc):
        """
        Rellena la tabla de puntos geodésicos SOLO con los puntos válidos del Excel.
        Elimina cualquier fila que contenga placeholders no reemplazados (como {{PUNTO_4}}).
        """
        if self.puntos_geodesicos.empty:
            return

        for table in doc.tables:
            if not table.rows or len(table.rows[0].cells) < 2:
                continue

            header_text = table.cell(0, 1).text.strip().upper()
            if 'CÓDIGOS DE PUNTOS GEODÉSICOS' in header_text or 'CODIGOS DE PUNTOS GEODESICOS' in header_text:
                if len(table.rows) < 2:
                    continue
                # Guardar la estructura de la fila plantilla (segunda fila)
                template_row = table.rows[1]
                tbl = table._tbl
                # Eliminar TODAS las filas de datos (excepto la cabecera)
                while len(table.rows) > 1:
                    row = table.rows[1]
                    row._element.getparent().remove(row._element)

                # Agregar una fila por cada punto válido
                for idx, punto_row in enumerate(self.puntos_geodesicos.itertuples(index=False), 1):
                    new_tr = deepcopy(template_row._tr)
                    new_row = _Row(new_tr, table)
                    for col_idx in range(len(new_row.cells)):
                        header = table.rows[0].cells[col_idx].text.strip().upper()
                        if col_idx == 0:
                            new_row.cells[0].text = str(idx)
                        else:
                            val = ''
                            for col in self.puntos_geodesicos.columns:
                                if header in str(col).strip().upper() or str(col).strip().upper() in header:
                                    val = getattr(punto_row, col)
                                    break
                            new_row.cells[col_idx].text = str(val) if pd.notna(val) else ''
                    tbl.append(new_tr)

                # Eliminar cualquier fila que contenga un placeholder no reemplazado
                filas_a_eliminar = []
                for i, row in enumerate(table.rows[1:], 1): 
                    for cell in row.cells:
                        if "{{" in cell.text or "}}" in cell.text:
                            filas_a_eliminar.append(i)
                            break
                for i in reversed(filas_a_eliminar):
                    row = table.rows[i]
                    row._element.getparent().remove(row._element)
                break

    def generar_formularios(self, image_paths_dict=None, zona=None, orden=None):
        """
        Genera los formularios Word seleccionados.
        image_paths_dict: dict opcional. Si se provee, debe ser un dict con clave el nombre del punto (str, igual que en la columna 'PUNTO'),
        y valor un dict con las rutas de imagen para ese punto, por ejemplo:
        {
            'PUNTO1': {
                'RUTA_ALTURA_DE_ANTENA': 'ruta/a/imagen1.jpg',
                'RUTA_UBICACION_PG': 'ruta/a/imagen2.jpg',
                ...
            },
            ...
        }
        Si no se provee, se usará el método anterior (DataFrame).
        """
        expediente_nombre = os.path.splitext(os.path.basename(self.excel_path))[0]
        zip_filename = os.path.join(self.output_dir, f"FORMULARIOS_{expediente_nombre}.zip")

        forms_per_point = ["Formulario 004", "Formulario 005"]

        with tempfile.TemporaryDirectory() as temp_dir:
            for form_name in self.seleccionados:
                template_path = resource_path(os.path.join("Assets", "Templates", f"{form_name.lower().replace(' ', '_')}.docx"))
                if not os.path.exists(template_path):
                    continue

                if form_name == "Formulario 001":
                    # --- Solo hoja 1 ---
                    doc = Document(template_path)
                    contexto_general = self._contexto_formulario_001()
                    self._fill_geodetic_points_table(doc)
                    self._replace_all_placeholders(doc, contexto_general)
                    # --- Insertar imagen de firma si existe ---
                    placeholders_imagenes = {}
                    firma_path = contexto_general.get('RUTA_FIRMA', '')
                    if firma_path and isinstance(firma_path, str) and os.path.exists(firma_path):
                        placeholders_imagenes['{{FIRMA}}'] = firma_path
                    else:
                        placeholders_imagenes['{{FIRMA}}'] = ''
                    self.reemplazar_imagenes_en_docx(doc, placeholders_imagenes)
                    doc.save(os.path.join(temp_dir, f"{form_name}.docx"))
                elif form_name in ["Formulario 002", "Formulario 003"]:
                    # --- Solo hoja 2 ---
                    doc = Document(template_path)
                    contexto_23 = {}
                    if hasattr(self, 'df_form2_3') and not self.df_form2_3.empty:
                        contexto_23 = self.df_form2_3.iloc[0].to_dict()
                    contexto_general = {key: str(value) for key, value in contexto_23.items()}
                    if form_name == "Formulario 003":
                        self._fill_geodetic_points_table(doc)
                    self._replace_all_placeholders(doc, contexto_general)
                    # --- Insertar imagen de firma si existe ---
                    placeholders_imagenes = {}
                    firma_path = contexto_general.get('RUTA_FIRMA', '')
                    if firma_path and isinstance(firma_path, str) and os.path.exists(firma_path):
                        placeholders_imagenes['{{FIRMA}}'] = firma_path
                    else:
                        placeholders_imagenes['{{FIRMA}}'] = ''
                    self.reemplazar_imagenes_en_docx(doc, placeholders_imagenes)
                    doc.save(os.path.join(temp_dir, f"{form_name}.docx"))

                elif form_name in ["Formulario 004", "Formulario 005"]:
                    # --- Hoja 3: Generar un formulario por cada punto geodésico válido ---
                    if self.puntos_geodesicos.empty:
                        # print(f"Advertencia: No hay puntos geodésicos para generar {form_name}.")
                        continue

                    # Leer datos generales de la hoja 3: hasta la fila donde empieza la tabla de puntos geodésicos
                    hoja3_raw = pd.read_excel(self.excel_path, sheet_name=2, header=None)
                    datos_generales_hoja3 = {}
                    def normaliza_clave(clave):
                        clave = str(clave).strip()
                        clave = unicodedata.normalize('NFKD', clave).encode('ASCII', 'ignore').decode('ASCII')
                        clave = clave.replace(' ', '_').replace('-', '_').replace('__', '_').upper()
                        return clave
                    # Tomar filas 2 a 14 (índices 1 a 13) como datos generales

                    for i in range(1, 14):
                        if i >= len(hoja3_raw):
                            break
                        clave_original = str(hoja3_raw.iloc[i, 0]).strip()
                        valor = hoja3_raw.iloc[i, 1] if hoja3_raw.shape[1] > 1 else ''
                        if clave_original != '':
                            clave_norm = normaliza_clave(clave_original)
                            # Omitir FIRMA y RUTA_FIRMA completamente
                            if clave_norm in ['FIRMA', 'RUTA_FIRMA']:
                                continue
                            valor_str = str(valor).strip() if pd.notna(valor) else ''
                            # Si la clave es ZONA, ZONA_UTM u ORDEN y se recibió un valor externo, usarlo
                            if zona is not None and (clave_norm == 'ZONA' or clave_norm == 'ZONA_UTM'):
                                datos_generales_hoja3[clave_norm] = str(zona)
                            elif clave_norm == 'ORDEN' and orden is not None:
                                datos_generales_hoja3[clave_norm] = str(orden)
                            else:
                                datos_generales_hoja3[clave_norm] = valor_str

                    def normaliza_clave(clave):
                        clave = str(clave).strip()
                        clave = unicodedata.normalize('NFKD', clave).encode('ASCII', 'ignore').decode('ASCII')
                        clave = clave.replace(' ', '_').replace('-', '_').replace('__', '_').upper()
                        return clave


                    # Variable para "recordar" el tipo de medida del primer punto válido
                    tipo_medida_default = None
                    for index, row in self.puntos_geodesicos.iterrows():
                        punto_val = row.get('PUNTO')
                        if pd.isna(punto_val) or str(punto_val).strip() == '':
                            continue
                        punto_nombre = str(punto_val).replace(" ", "_")
                        # Copiar datos generales y luego sobreescribir con los del punto
                        contexto_para_este_punto = {key: value for key, value in datos_generales_hoja3.items()}
                        # Sobrescribir ZONA, ZONA_UTM y ORDEN con los valores externos si se recibieron
                        if zona is not None:
                            contexto_para_este_punto['ZONA'] = str(zona)
                            contexto_para_este_punto['ZONA_UTM'] = str(zona)
                        if orden is not None:
                            contexto_para_este_punto['ORDEN'] = str(orden)
                        # También agregar variantes normalizadas de ZONA_UTM si zona viene del frontend
                        if zona is not None:
                            def normaliza_clave(clave):
                                clave = str(clave).strip()
                                clave = unicodedata.normalize('NFKD', clave).encode('ASCII', 'ignore').decode('ASCII')
                                clave = clave.replace(' ', '_').replace('-', '_').replace('__', '_').upper()
                                return clave
                            variantes_zona_utm = set()
                            variantes_zona_utm.add('ZONA_UTM')
                            variantes_zona_utm.add(normaliza_clave('ZONA_UTM'))
                            variantes_zona_utm.add('ZONA UTM')
                            variantes_zona_utm.add(normaliza_clave('ZONA UTM'))
                            variantes_zona_utm.add('ZONAUTM')
                            variantes_zona_utm.add(normaliza_clave('ZONAUTM'))
                            for variante in variantes_zona_utm:
                                contexto_para_este_punto[variante] = str(zona)
                        for col in row.index:
                            col_str = str(col).upper()
                            # Si la clave es ZONA u ORDEN y ya fue establecida por el frontend, no sobrescribir
                            if (col_str == 'ZONA' and zona is not None) or (col_str == 'ORDEN' and orden is not None):
                                continue
                            valor_col = self._safe_get(row, col)
                            # Solo sobrescribir si el valor no está vacío
                            if valor_col is not None and str(valor_col).strip() != '':
                                contexto_para_este_punto[col_str] = valor_col

                        contexto_para_este_punto['COD_PUNTO _GEODESICO'] = contexto_para_este_punto.get('PUNTO', f'sin_nombre_{index}')

                        # --- LÓGICA CON MEMORIA DEL PRIMER PUNTO ---
                        # 1. Intentar leer el valor para el punto actual
                        tipo_medida_actual = None
                        possible_keys = ['TIPO_DE_MEDIDA', 'TIPO DE MEDIDA']
                        for key in possible_keys:
                            if key in row.index:
                                val = row.get(key)
                                if pd.notna(val) and str(val).strip() != '':
                                    tipo_medida_actual = str(val).strip()
                                    break
                        
                        # 2. Si se encontró un valor, usarlo y guardarlo como default si es el primero.
                        if tipo_medida_actual:
                            tipo_a_usar = tipo_medida_actual
                            if tipo_medida_default is None:
                                tipo_medida_default = tipo_medida_actual
                        # 3. Si no se encontró valor, usar el default guardado.
                        else:
                            tipo_a_usar = tipo_medida_default

                        # 4. Determinar V e I basados en el tipo a usar (actual o default)
                        v_val, i_val = '', ''
                        if tipo_a_usar:
                            tipo_medida = str(tipo_a_usar).strip().lower()
                            tipo_medida = tipo_medida.replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
                            tipo_medida = ' '.join(tipo_medida.split())
                            
                            if 'vertical' in tipo_medida:
                                v_val = 'X'
                            elif 'inclinada' in tipo_medida:
                                i_val = 'X'

                        contexto_para_este_punto['V'] = v_val
                        contexto_para_este_punto['I'] = i_val
                        # --- Fin de la lógica con memoria ---

                        # Generar variantes normalizadas de las claves para asegurar reemplazo de placeholders
                        contexto_extendido = {}
                        # Primero, agregar los datos generales con variantes normalizadas
                        for k, v in datos_generales_hoja3.items():
                            claves = set()
                            claves.add(k)
                            claves.add(k.upper())
                            claves.add(normaliza_clave(k))
                            claves.add(normaliza_clave(k.lower()))
                            claves.add(normaliza_clave(k.upper()))
                            k_sin_paren = re.sub(r'[\(\)]', '', k)
                            claves.add(k_sin_paren)
                            claves.add(normaliza_clave(k_sin_paren))
                            k_sin_guion = k.replace('-', '_')
                            claves.add(k_sin_guion)
                            claves.add(normaliza_clave(k_sin_guion))
                            k_sin_espacios = k.replace(' ', '')
                            claves.add(k_sin_espacios)
                            claves.add(normaliza_clave(k_sin_espacios))
                            for variante in claves:
                                contexto_extendido[variante] = v
                        # Luego, agregar los datos del punto (sobrescriben si hay coincidencia, excepto V/I)
                        for k, v in contexto_para_este_punto.items():
                            clave_norm = k.strip('{}').replace(' ', '_').replace('-', '_').upper()
                            if clave_norm in ['V', 'I']:
                                continue  
                            
                            claves = set()
                            claves.add(k)
                            claves.add(k.upper())
                            claves.add(normaliza_clave(k))
                            claves.add(normaliza_clave(k.lower()))
                            claves.add(normaliza_clave(k.upper()))
                            k_sin_paren = re.sub(r'[\(\)]', '', k)
                            claves.add(k_sin_paren)
                            claves.add(normaliza_clave(k_sin_paren))
                            k_sin_guion = k.replace('-', '_')
                            claves.add(k_sin_guion)
                            claves.add(normaliza_clave(k_sin_guion))
                            k_sin_espacios = k.replace(' ', '')
                            claves.add(k_sin_espacios)
                            claves.add(normaliza_clave(k_sin_espacios))
                            for variante in claves:
                                clave_var_norm = variante.strip('{}').replace(' ', '_').replace('-', '_').upper()
                                if clave_var_norm in ['FIRMA', 'RUTA_FIRMA']:
                                    # Si el valor del punto NO está vacío, sobrescribir
                                    if v is not None and str(v).strip() != '':
                                        contexto_extendido[variante] = v
                                    # Si el valor del punto está vacío, asegurar que la clave siga en el contexto con el valor general
                                    elif variante not in contexto_extendido:
                                        contexto_extendido[variante] = datos_generales_hoja3.get(variante, '')
                                else:
                                    if v is not None and str(v).strip() != '':
                                        contexto_extendido[variante] = v

                        # Finalmente, asegurar variantes normalizadas de V e I en el contexto_extendido (después de todo lo anterior)
                        for clave, valor in {'V': v_val, 'I': i_val}.items():
                            variantes = set()
                            variantes.add(clave)
                            variantes.add(clave.upper())
                            variantes.add(normaliza_clave(clave))
                            variantes.add(normaliza_clave(clave.lower()))
                            variantes.add(normaliza_clave(clave.upper()))
                            k_sin_paren = re.sub(r'[\(\)]', '', clave)
                            variantes.add(k_sin_paren)
                            variantes.add(normaliza_clave(k_sin_paren))
                            k_sin_guion = clave.replace('-', '_')
                            variantes.add(k_sin_guion)
                            variantes.add(normaliza_clave(k_sin_guion))
                            k_sin_espacios = clave.replace(' ', '')
                            variantes.add(k_sin_espacios)
                            variantes.add(normaliza_clave(k_sin_espacios))
                            # Agregar variantes con llaves
                            variantes.add(f'{{{{{clave}}}}}')
                            variantes.add(f'{{{{{clave.upper()}}}}}')
                            variantes.add(f'{{{{{normaliza_clave(clave)}}}}}')
                            variantes.add(f'{{{{{normaliza_clave(clave.lower())}}}}}')
                            variantes.add(f'{{{{{normaliza_clave(clave.upper())}}}}}')
                            variantes.add(f'{{{{{k_sin_paren}}}}}')
                            variantes.add(f'{{{{{normaliza_clave(k_sin_paren)}}}}}')
                            variantes.add(f'{{{{{k_sin_guion}}}}}')
                            variantes.add(f'{{{{{normaliza_clave(k_sin_guion)}}}}}')
                            variantes.add(f'{{{{{k_sin_espacios}}}}}')
                            variantes.add(f'{{{{{normaliza_clave(k_sin_espacios)}}}}}')
                            for variante in variantes:
                                contexto_extendido[variante] = valor if valor is not None else ''


                        # Depuración: mostrar el valor de todas las variantes de V e I en el contexto
                        variantes_v = [k for k in contexto_extendido.keys() if 'V' in k.upper() and len(k.replace('{','').replace('}','')) <= 3]
                        variantes_i = [k for k in contexto_extendido.keys() if 'I' in k.upper() and len(k.replace('{','').replace('}','')) <= 3]

                        doc_punto = Document(template_path)
                        self._replace_all_placeholders(doc_punto, contexto_extendido)



                        # Usar siempre las rutas de imagenes cargadas al programa (image_paths_dict), si existen
                        img_dict = image_paths_dict.get(punto_nombre, {}) if image_paths_dict else {}
                        placeholders_imagenes = {
                            "{{IMAGEN_MEDICION_ALTURA_DE_LA_ANTENA}}": img_dict.get("RUTA_ALTURA_DE_ANTENA") or self._safe_get(row, "RUTA_ALTURA_DE_ANTENA"),
                            "{{REF_UBICACION_PG}}": img_dict.get("RUTA_UBICACION_PG") or self._safe_get(row, "RUTA_UBICACION_PG"),
                            "{{IMAGEN_POSICIONAMIENTO_GPS_GNSS}}": img_dict.get("RUTA_POSICIONAMIENTO_GPS_GNSS") or self._safe_get(row, "RUTA_POSICIONAMIENTO_GPS_GNSS"),
                            "{{IMAGEN_DISCO_DE_BRONCE_POSICIONAMIENTO}}": img_dict.get("RUTA_DISCO_DE_BRONCE") or self._safe_get(row, "RUTA_DISCO_DE_BRONCE")
                        }

                        self.reemplazar_imagenes_en_docx(doc_punto, placeholders_imagenes)

                        output_filename = f"{form_name}_{punto_nombre}.docx"
                        doc_punto.save(os.path.join(temp_dir, output_filename))

            if not os.listdir(temp_dir):
                raise FileNotFoundError("No se generó ningún formulario. Verifique los datos del Excel y que las plantillas existan.")
            
            shutil.make_archive(os.path.splitext(zip_filename)[0], 'zip', temp_dir)
        
        return zip_filename

    def reemplazar_imagenes_en_docx(self, doc, placeholders_imagenes):
        """
        Reemplaza los placeholders de imagen en un documento docx por la ruta de la imagen como texto (NO inserta la imagen, solo la ruta),
        excepto para {{FIRMA}}, donde inserta la imagen si la ruta existe.
        Solo inserta una vez la imagen por párrafo/celda aunque el placeholder esté repetido.
        """

        # Definir dimensiones personalizadas para cada placeholder
        placeholder_dims = {
            "{{IMAGEN_MEDICION_ALTURA_DE_LA_ANTENA}}": {"width": 8.5, "height": 10.16},
            "{{REF_UBICACION_PG}}": {"width": 9.7, "height": 10.1},
            "{{IMAGEN_POSICIONAMIENTO_GPS_GNSS}}": {"width": 5.56, "height": 4.56},
            "{{IMAGEN_DISCO_DE_BRONCE_POSICIONAMIENTO}}": {"width": 5.56, "height": 4.56},
        }


        # Reemplazo en párrafos principales
        for p in doc.paragraphs:
            for key, img_path in placeholders_imagenes.items():
                if key in p.text:
                    full_text = ''.join(run.text for run in p.runs)
                    if key in full_text:
                        partes = full_text.split(key)
                        for run in p.runs[:]:
                            p._element.remove(run._element)
                        for i, parte in enumerate(partes):
                            if parte:
                                p.add_run(parte)
                            if i < len(partes) - 1:
                                if img_path and isinstance(img_path, str) and os.path.exists(img_path):
                                    dims = placeholder_dims.get(key)
                                    if dims:
                                        run = p.add_run()
                                        run.add_picture(img_path, width=Cm(dims["width"]), height=Cm(dims["height"]))
                                    else:
                                        p.add_run().add_picture(img_path, width=Inches(1.5))
                                elif img_path and isinstance(img_path, str):
                                    p.add_run(img_path)

        # Reemplazo en tablas (todas las celdas)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, img_path in placeholders_imagenes.items():
                            if key in p.text:
                                full_text = ''.join(run.text for run in p.runs)
                                if key in full_text:
                                    partes = full_text.split(key)
                                    for run in p.runs[:]:
                                        p._element.remove(run._element)
                                    for i, parte in enumerate(partes):
                                        if parte:
                                            p.add_run(parte)
                                        if i < len(partes) - 1:
                                            if img_path and isinstance(img_path, str) and os.path.exists(img_path):
                                                dims = placeholder_dims.get(key)
                                                if dims:
                                                    run = p.add_run()
                                                    run.add_picture(img_path, width=Cm(dims["width"]), height=Cm(dims["height"]))
                                                else:
                                                    p.add_run().add_picture(img_path, width=Inches(1.5))
                                            elif img_path and isinstance(img_path, str):
                                                p.add_run(img_path)
        return doc